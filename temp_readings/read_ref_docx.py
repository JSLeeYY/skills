# -*- coding: utf-8 -*-
from docx import Document

doc = Document(r'D:\Desktop' + '\\' + '\u848b\u6b63\u70e8\u6570\u636e\u5e93\u88682.1.docx')

for i, para in enumerate(doc.paragraphs):
    style_name = para.style.name if para.style else 'None'
    text = para.text
    if text.strip():
        print('P{} [{}]: {}'.format(i, style_name, text))

print('\n=== Tables ===')
for t_idx, table in enumerate(doc.tables):
    print('\nTable {}:'.format(t_idx + 1))
    for r_idx, row in enumerate(table.rows):
        cells = [cell.text for cell in row.cells]
        print('  Row {}: {}'.format(r_idx, ' | '.join(cells)))
