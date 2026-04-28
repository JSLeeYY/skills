# -*- coding: utf-8 -*-
import json
import os
import io
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Configuration
SOURCE_JSON = r"D:\Desktop\三方平台项目\三期需求\extraction.json"
OUTPUT_DOCX = r"D:\Desktop\三方平台项目\三期需求\数字一体化平台需求完整分析.docx"

# System Mapping (Ensure unicode literals for Py2)
SYSTEM_MAPPING = {
    u"执行管理与经营分析需求表.xlsx": u"执行管理与经营分析系统",
    u"预评价需求.xlsx": u"PECMS系统 - 预评价模块",
    u"后评价需求.xlsx": u"PECMS系统 - 后评价模块",
    u"数字信息平台优化提升需求评估 - 可读.xlsx": u"数字信息平台优化提升",
}

def set_font(run, font_name=u'微软雅黑', size=None, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size:
        run.font.size = Pt(size)
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level):
    h = doc.add_heading(level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    run.text = text
    set_font(run, size=16 if level==1 else 14, bold=True, color=RGBColor(0, 51, 102))

def create_table_from_rows(doc, rows_data):
    if not rows_data:
        return
    
    # Analyze columns
    parsed_rows = [row.split('|') for row in rows_data]
    max_cols = max(len(r) for r in parsed_rows) if parsed_rows else 0
    
    if max_cols == 0:
        return

    table = doc.add_table(rows=len(parsed_rows), cols=max_cols)
    table.style = 'Table Grid'
    
    for i, row_cells in enumerate(parsed_rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_cells):
            if j < len(row.cells):
                cell = row.cells[j]
                # python-docx in py2 might return something different, but usually ok
                if not cell.paragraphs:
                    p = cell.add_paragraph()
                else:
                    p = cell.paragraphs[0]
                
                # Ensure cell_text is unicode
                if not isinstance(cell_text, unicode):
                    cell_text = unicode(cell_text, 'utf-8', 'ignore') if hasattr(cell_text, 'decode') else unicode(cell_text)

                run = p.add_run(cell_text)
                set_font(run, size=10)
                # Header row formatting
                if i == 0:
                    set_font(run, size=10, bold=True)

def process_excel(doc, filename, data):
    for sheet_name, content in data.items():
        if isinstance(content, list):
            doc.add_heading(u"{}".format(sheet_name), level=2)
            create_table_from_rows(doc, content)
            doc.add_paragraph(u"")

def process_word_report(doc, filename, data):
    # Text content
    text = data.get("Text", u"")
    if not isinstance(text, unicode):
        text = unicode(text, 'utf-8', 'ignore') if hasattr(text, 'decode') else unicode(text)
        
    lines = [l.strip() for l in text.split('\r')]
    cleaned_lines = [l for l in lines if l]
    
    # Add Text
    for line in cleaned_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_font(run, size=11)
        
    # Comments
    comments = data.get("Comments", [])
    if comments:
        doc.add_heading(u"批注/修改意见", level=3)
        for c in comments:
            author = c.get("A", "Unknown")
            cmt_text = c.get("T", "")
            
            # Unicode safety
            if not isinstance(author, unicode): author = unicode(author, 'utf-8', 'ignore') if hasattr(author, 'decode') else unicode(author)
            if not isinstance(cmt_text, unicode): cmt_text = unicode(cmt_text, 'utf-8', 'ignore') if hasattr(cmt_text, 'decode') else unicode(cmt_text)

            p = doc.add_paragraph()
            run_auth = p.add_run(u"【{}】: ".format(author))
            set_font(run_auth, bold=True, color=RGBColor(255, 0, 0))
            run_txt = p.add_run(cmt_text)
            set_font(run_txt, color=RGBColor(255, 0, 0))

def main():
    if not os.path.exists(SOURCE_JSON):
        print("Error: JSON not found.")
        return

    # Use io.open for reliable utf-8 reading in Py2
    with io.open(SOURCE_JSON, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()
    
    # Title
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(u"数字一体化平台需求完整分析")
    set_font(run, size=24, bold=True, color=RGBColor(0, 0, 0))

    if "Status" in data:
        del data["Status"]

    processed_files = set()
    
    # 1. Excel Requirements
    for filename, system_name in SYSTEM_MAPPING.items():
        # filename is unicode from mapping
        if filename in data:
            add_heading(doc, system_name, level=1)
            process_excel(doc, filename, data[filename])
            processed_files.add(filename)

    # 2. Reports
    doc.add_page_break()
    add_heading(doc, u"现有报告分析与批注", level=1)
    
    for filename, content in data.items():
        if filename in processed_files:
            continue
        
        # filename from json load is unicode
        
        if isinstance(content, dict) and ("Text" in content or "Comments" in content):
            add_heading(doc, filename, level=2)
            process_word_report(doc, filename, content)
        elif isinstance(content, dict): 
            add_heading(doc, u"其他文件: {}".format(filename), level=2)
            process_excel(doc, filename, content)
            
    doc.save(OUTPUT_DOCX)
    print("Successfully generated docx.")

if __name__ == "__main__":
    main()
