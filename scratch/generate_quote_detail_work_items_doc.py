from __future__ import annotations

import copy
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook


BASE_DIR = Path(
    r"D:\DevelopmentLocation\agent skill\skills\word_project\2026\2026-04-24\01-谈判文件对外发布版本内容编辑"
)
SOURCE_XLSX = BASE_DIR / "20260413-三期需求汇总表-评估(1)(2).xlsx"
SOURCE_DOCX = BASE_DIR / "南京三方一体化数字信息平台优化提升项目谈判文件(2025-新增供应商考核)(修订稿).docx"
OUTPUT_DOCX = BASE_DIR / "分项报价表-工作内容费用项替换稿.docx"


def norm(value: object) -> str:
    if value is None:
        return ""
    return str(value).replace("\r\n", "\n").replace("\xa0", " ").strip()


def set_run_font(run, size: float = 10.5, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def set_cell_text(cell, text: str, *, center: bool = False) -> None:
    cell.text = ""
    pieces = text.split("\n") if text else [""]
    for idx, piece in enumerate(pieces):
        para = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        run = para.add_run(piece)
        set_run_font(run, size=10.5)


def append_grouped_text(lines: list[str]) -> str:
    cleaned = [line for line in lines if line]
    return "\n".join(cleaned)


def gather_pecms_items(ws, scope_keyword: str) -> OrderedDict[str, list[str]]:
    data: OrderedDict[str, list[str]] = OrderedDict()
    current_module = ""
    current_scope = ""

    for row in ws.iter_rows(min_row=6, values_only=True):
        module = norm(row[0])
        detail = norm(row[1])
        scope = norm(row[11])
        if module in {"各项合计", "工作量汇总"}:
            break
        if module.startswith("附：") or module in {"开放接口方案一", "开放接口方案二"}:
            continue
        if module:
            current_module = module
        if scope:
            current_scope = scope
        if not detail:
            continue
        if scope_keyword not in current_scope:
            continue
        data.setdefault(current_module, []).append(detail)

    return data


def gather_app_items(ws) -> OrderedDict[str, list[str]]:
    data: OrderedDict[str, list[str]] = OrderedDict()
    current_domain = ""
    current_module = ""
    for row in ws.iter_rows(min_row=6, values_only=True):
        domain = norm(row[0])
        module = norm(row[1])
        detail = norm(row[2])
        if module in {"各项合计", "工作量汇总"}:
            break
        if domain:
            current_domain = domain
        if module:
            current_module = module
        if not detail:
            continue
        title = f"{current_domain}-{current_module}" if current_domain and current_module else current_module or current_domain
        data.setdefault(title, []).append(detail)
    return data


def gather_safety_items(ws) -> OrderedDict[str, list[str]]:
    data: OrderedDict[str, list[str]] = OrderedDict()
    current_module = ""
    for row in ws.iter_rows(min_row=6, values_only=True):
        module = norm(row[0])
        detail = norm(row[1])
        remark = norm(row[7])
        supplement = norm(row[9])
        if module in {"各项合计", "工作量汇总"}:
            break
        if module:
            current_module = module
        parts = [detail]
        if remark:
            parts.append(f"备注：{remark}")
        if supplement:
            parts.append(f"补充说明：{supplement}")
        merged = " ".join(part for part in parts if part)
        if merged:
            data.setdefault(current_module, []).append(merged)
    return data


def gather_om_items(ws) -> OrderedDict[str, list[str]]:
    data: OrderedDict[str, list[str]] = OrderedDict()
    current_module = ""
    for row in ws.iter_rows(min_row=6, values_only=True):
        module = norm(row[0])
        detail = norm(row[1])
        remark = norm(row[7])
        if module in {"各项合计", "工作量汇总"}:
            break
        if module:
            current_module = module
        if not detail:
            continue
        merged = detail if not remark else f"{detail} 备注：{remark}"
        data.setdefault(current_module, []).append(merged)
    return data


def render_section(title: str, grouped: OrderedDict[str, list[str]]) -> str:
    lines = [f"{title}："]
    idx = 1
    for module, details in grouped.items():
        subparts = []
        for sub_idx, detail in enumerate(details, 1):
            clean_detail = detail.replace("\n", " ").strip()
            if clean_detail:
                subparts.append(f"（{sub_idx}）{clean_detail}")
        if subparts:
            lines.append(f"{idx}. {module}：{'；'.join(subparts)}。")
            idx += 1
    return append_grouped_text(lines)


def build_work_item_texts() -> list[str]:
    wb = load_workbook(SOURCE_XLSX, data_only=True)

    pecms_ws = wb["PECMS(对内+对外PC)"]
    app_ws = wb["APP"]
    safety_ws = wb["安全管理"]
    om_ws = wb["经营管理分析"]

    pecms_inner = gather_pecms_items(pecms_ws, "对内")
    pecms_outer = gather_pecms_items(pecms_ws, "对外")

    # 对外开放接口单独补入，避免仅靠分类字段时遗漏方案说明
    outer_interface = [
        "开放接口方案开发及联调（以最终确认方案为准），包含检验主管查询、业主基础信息及附件接收，以及按项目/派单维度对文档目录、历史版本、细则必选表单、交底、人员交接、到岗、设备申请、日/月报、监理大纲、日志、监理实施细则、监理报告、监理总结等数据的标准接口查询与交互能力建设。"
    ]
    pecms_outer.setdefault("开放接口方案", []).extend(outer_interface)

    # 对内侧将与技术服务、合同、消息、项目状态相关的对内专属能力保留
    inner_text = render_section("对内PECMS平台开发", pecms_inner)
    outer_text = render_section("对外PECMS平台开发", pecms_outer)
    app_text = render_section("对内PECMS-APP端功能优化改造", gather_app_items(app_ws))
    om_text = render_section("执行管理与经营分析平台开发", gather_om_items(om_ws))
    safety_text = render_section("安全管理平台升级改造", gather_safety_items(safety_ws))

    return [inner_text, outer_text, app_text, om_text, safety_text]


def clone_quote_detail_table() -> Document:
    src = Document(SOURCE_DOCX)
    table = src.tables[9]

    doc = Document()
    # 统一基础字体
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = Pt(10.5)

    cloned_tbl = copy.deepcopy(table._tbl)
    doc._body._element.append(cloned_tbl)
    new_table = doc.tables[0]

    # 删除上半部分“统一报价表”，只保留“报价明细表”标题+表头+5行明细
    for _ in range(10):
        new_table._tbl.remove(new_table._tbl.tr_lst[0])

    return doc


def fill_table(doc: Document, work_texts: list[str]) -> None:
    table = doc.tables[0]
    # row0: merged title row; row1: header; row2-6 data
    for idx, text in enumerate(work_texts, start=2):
        row = table.rows[idx]
        # 序号保留，其他列清空，仅填“工作内容/费用项”
        for col_idx, cell in enumerate(row.cells):
            if col_idx == 0:
                set_cell_text(cell, str(idx - 1), center=True)
            elif col_idx == 1:
                set_cell_text(cell, text, center=False)
            else:
                set_cell_text(cell, "", center=False)


def main() -> None:
    work_texts = build_work_item_texts()
    doc = clone_quote_detail_table()
    fill_table(doc, work_texts)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
