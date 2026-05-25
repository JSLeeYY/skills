from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
HEADER_FILL = "F2F4F7"
ACCENT_BLUE = RGBColor(0x2E, 0x74, 0xB5)
ACCENT_DARK = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)
BLACK = RGBColor(0x00, 0x00, 0x00)
CN_FONT = "Microsoft YaHei"
LATIN_FONT = "Calibri"


@dataclass
class InterfaceItem:
    index: int
    name: str
    if_id: str
    method: str
    uri: str


def set_run_font(run, size: float | None = None, *, bold: bool = False, color: RGBColor | None = None) -> None:
    run.bold = bold
    run.font.name = LATIN_FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_format(paragraph, *, before: float = 0, after: float = 6, line: float = 1.1, align=None) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = tc_mar.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def set_table_layout_fixed(table, col_widths_dxa: Iterable[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in col_widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, col_widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def repeat_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def apply_table_style(table, header_fill: str = HEADER_FILL) -> None:
    table.style = "Table Grid"
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for para in cell.paragraphs:
                set_paragraph_format(para, before=0, after=0, line=1.15)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in (0, 2, 3, 5) else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    set_run_font(run, size=10.5, bold=(r_idx == 0))
            if r_idx == 0:
                shade_cell(cell, header_fill)
    repeat_header_row(table.rows[0])


def add_cell_text(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT, size: float = 10.5, color=BLACK) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_format(p, before=0, after=0, line=1.15, align=align)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def set_page_geometry(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    normal.font.size = Pt(11)


def set_header_footer(section, company_name: str) -> None:
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    set_paragraph_format(hp, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = hp.add_run("PECMS接口验收报告")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    set_paragraph_format(fp, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = fp.add_run(f"适用对象：{company_name}")
    set_run_font(run, size=9, color=MUTED)


def add_title_block(doc: Document, company_name: str, interface_total: int, method_counts: dict[str, int], base_url: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, before=24, after=4, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("PECMS接口验收报告")
    set_run_font(run, size=20, bold=True, color=BLACK)

    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=18, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(f"五环工程公司接口交付验收文件")
    set_run_font(run, size=11.5, color=MUTED)

    meta_rows = [
        ("甲方", company_name),
        ("乙方", "[待填写]"),
        ("交付内容", "PECMS公共接口"),
        ("接口总数", f"{interface_total} 项"),
        ("请求地址前缀", base_url or "https://pecms.com.cn"),
        ("请求方式统计", "，".join(f"{k} {v} 项" for k, v in method_counts.items() if k) or "[待填写]"),
        ("验收日期", "[待填写]"),
        ("验收结论", "通过"),
    ]

    table = doc.add_table(rows=1, cols=2)
    for label, value in meta_rows:
        row = table.rows[0] if len(table.rows) == 1 and table.rows[0].cells[0].text == "" and table.rows[0].cells[1].text == "" else table.add_row()
        add_cell_text(row.cells[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, color=ACCENT_DARK)
        add_cell_text(row.cells[1], value, align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5)
    set_table_layout_fixed(table, [1800, 7560])
    apply_table_style(table)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    if level == 1:
        set_paragraph_format(p, before=16, after=8, line=1.0)
        run = p.add_run(text)
        set_run_font(run, size=16, bold=True, color=ACCENT_BLUE)
    elif level == 2:
        set_paragraph_format(p, before=12, after=6, line=1.0)
        run = p.add_run(text)
        set_run_font(run, size=13, bold=True, color=ACCENT_BLUE)
    else:
        set_paragraph_format(p, before=8, after=4, line=1.0)
        run = p.add_run(text)
        set_run_font(run, size=12, bold=True, color=ACCENT_DARK)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=6, line=1.1)
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)


def add_numbered_line(doc: Document, index: int, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=4, line=1.15)
    run = p.add_run(f"{index}. {text}")
    set_run_font(run, size=11, color=BLACK)


def add_summary_table(doc: Document) -> None:
    rows = [
        ("接口范围", "本次验收仅针对乙方向甲方提供的 PECMS 公共接口。", "通过", "未纳入平台部署、服务器、数据库、源码及运维服务内容。"),
        ("清单对应性", "接口名称、if-id、请求方式、请求 uri 与《PECMS接口清单-五环工程公司》一致。", "通过", "以接口清单为验收范围边界。"),
        ("鉴权规则", "接口请求头已明确 access-key、timestamp、if-id、sign 等调用规则。", "通过", "满足接口调用识别与签名校验说明。"),
        ("参数定义", "各接口已明确 Path、Body、Query 等请求参数及是否必填。", "通过", "接口调用入参定义完整。"),
        ("响应结构", "各接口已明确响应状态码、返回数据结构及示例报文。", "通过", "满足甲方接入和解析要求。"),
        ("文件能力", "涉及附件、图片、文档、报表等接口已提供访问说明。", "通过", "包含文件下载、图片预览及文档数据接口。"),
    ]

    table = doc.add_table(rows=1, cols=4)
    headers = ["序号", "验收项目", "验收结果", "说明"]
    for cell, text in zip(table.rows[0].cells, headers):
        add_cell_text(cell, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx, row_data in enumerate(rows, 1):
        row = table.add_row()
        add_cell_text(row.cells[0], str(idx), align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(row.cells[1], row_data[0])
        add_cell_text(row.cells[2], row_data[2], align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT_DARK)
        add_cell_text(row.cells[3], row_data[1] + row_data[3])
    set_table_layout_fixed(table, [720, 1800, 1080, 5760])
    apply_table_style(table)


def add_interface_table(doc: Document, items: list[InterfaceItem]) -> None:
    table = doc.add_table(rows=1, cols=6)
    headers = ["序号", "接口名称", "if-id", "方法", "请求 uri", "结果"]
    for cell, text in zip(table.rows[0].cells, headers):
        add_cell_text(cell, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for item in items:
        row = table.add_row()
        add_cell_text(row.cells[0], str(item.index), align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(row.cells[1], item.name)
        add_cell_text(row.cells[2], item.if_id, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(row.cells[3], item.method, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(row.cells[4], item.uri)
        add_cell_text(row.cells[5], "通过", align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT_DARK)
    set_table_layout_fixed(table, [600, 2100, 900, 900, 3600, 1260])
    apply_table_style(table)


def add_signature_table(doc: Document) -> None:
    table = doc.add_table(rows=4, cols=4)
    data = [
        ["单位", "代表/负责人", "签字", "日期"],
        ["甲方（五环工程公司）", "", "", ""],
        ["乙方", "", "", ""],
        ["备注", "本报告仅作为接口交付验收文件使用。", "", ""],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, text in enumerate(row_data):
            add_cell_text(
                table.rows[r_idx].cells[c_idx],
                text,
                bold=(r_idx == 0),
                align=WD_ALIGN_PARAGRAPH.CENTER if c_idx != 1 or r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT,
            )
    set_table_layout_fixed(table, [2200, 2200, 2480, 2480])
    apply_table_style(table)


def extract_base_url(doc: Document) -> str:
    for para in doc.paragraphs:
        text = para.text.strip()
        if "http" in text:
            match = re.search(r"https?://[^\s]+", text)
            if match:
                return match.group(0)
    return "https://pecms.com.cn"


def parse_interfaces(doc: Document) -> list[InterfaceItem]:
    texts = [p.text.strip() for p in doc.paragraphs]
    items: list[InterfaceItem] = []
    for idx, text in enumerate(texts):
        if "if-id" not in text:
            continue

        name = ""
        j = idx - 1
        while j >= 0:
            if texts[j]:
                name = texts[j]
                break
            j -= 1

        next_values: list[str] = []
        k = idx + 1
        while k < len(texts) and len(next_values) < 4:
            if texts[k]:
                next_values.append(texts[k])
            k += 1

        if_id_match = re.search(r"(\d+)", text)
        method = ""
        uri = ""
        for candidate in next_values:
            method_match = re.search(r"\b(GET|POST|PUT|DELETE|PATCH)\b", candidate, re.IGNORECASE)
            if method_match and not method:
                method = method_match.group(1).upper()
            if "/" in candidate and not uri:
                uri = candidate.split("：", 1)[-1].strip()

        items.append(
            InterfaceItem(
                index=len(items) + 1,
                name=name,
                if_id=if_id_match.group(1) if if_id_match else "",
                method=method,
                uri=uri,
            )
        )
    return items


def build_report(source_docx: Path, output_docx: Path) -> None:
    source = Document(str(source_docx))
    items = parse_interfaces(source)
    method_counts: dict[str, int] = {}
    for item in items:
        method_counts[item.method] = method_counts.get(item.method, 0) + 1

    company_name = "五环工程公司"
    base_url = extract_base_url(source)

    doc = Document()
    set_page_geometry(doc)
    set_header_footer(doc.sections[0], company_name)
    add_title_block(doc, company_name, len(items), method_counts, base_url)

    add_heading(doc, "一、验收说明", 1)
    add_body_paragraph(
        doc,
        "根据《PECMS接口清单-五环工程公司》，乙方向甲方提供 PECMS 公共接口服务。"
        "本报告仅针对接口交付内容进行验收，验收边界包括接口名称、接口编号（if-id）、请求方式、请求 uri、鉴权规则、请求参数、响应状态及响应结构；"
        "不包括平台部署、服务器环境、数据库、源代码、运维服务及其他非接口交付事项。"
    )

    add_heading(doc, "二、验收依据", 1)
    add_numbered_line(doc, 1, "《PECMS接口清单-五环工程公司.docx》。")
    add_numbered_line(doc, 2, "双方确认的接口对接范围、接口调用规则及接口报文说明。")

    add_heading(doc, "三、验收范围", 1)
    add_body_paragraph(
        doc,
        f"本次纳入验收的接口共 {len(items)} 项，请求地址前缀为 {base_url}。"
        f"其中 GET 接口 {method_counts.get('GET', 0)} 项，POST 接口 {method_counts.get('POST', 0)} 项。"
    )
    add_body_paragraph(
        doc,
        "接口清单覆盖项目、派单、设备、文件、图片、文档、报表、进度、异常问题、数据分析、监理日志、监理总结等对外提供的数据访问与查询能力。"
    )
    add_body_paragraph(
        doc,
        "接口调用请求头统一包含 access-key、timestamp、if-id、sign 等字段，满足甲方按统一规则进行调用、鉴权和接入的需要。"
    )

    add_heading(doc, "四、验收结果", 1)
    add_summary_table(doc)

    doc.add_page_break()
    add_heading(doc, "五、接口验收明细表", 1)
    add_body_paragraph(
        doc,
        "以下清单按照《PECMS接口清单-五环工程公司》逐项整理，作为本次接口交付验收明细。"
    )
    add_interface_table(doc, items)

    add_heading(doc, "六、验收结论", 1)
    add_body_paragraph(
        doc,
        "经对照《PECMS接口清单-五环工程公司》进行核对，本次提交的 PECMS 公共接口清单范围明确、接口调用规则清晰、请求方式与请求 uri 定义完整、"
        "请求参数与响应结构说明齐备，满足甲方对接口交付验收的要求。"
    )
    add_body_paragraph(
        doc,
        "综上，本次乙方向甲方提供的 PECMS 接口交付内容，同意通过接口验收。"
    )

    add_heading(doc, "七、签署确认", 1)
    add_signature_table(doc)

    doc.save(str(output_docx))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    build_report(args.source, args.output)
    print(args.output)


if __name__ == "__main__":
    main()
