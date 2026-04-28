from __future__ import annotations

import copy
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook


ROOT = Path.cwd()
DATE_DIR = ROOT / "word_project" / "2026" / "2026-04-24"
WORK_DIR = next(path for path in DATE_DIR.iterdir() if path.is_dir())
SOURCE_XLSX = next(path for path in WORK_DIR.glob("20260413*.xlsx"))
OUTPUT_DOCX = WORK_DIR / "分项报价表-按细分模块工作内容替换稿.docx"


def normalize(value: object) -> str:
    if value is None:
        return ""
    text = str(value).replace("\r\n", "\n").replace("\r", "\n").replace("\xa0", " ").strip()
    text = text.replace("\uf0b7", "- ")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def normalize_scope(scope: str) -> str:
    scope = normalize(scope)
    if scope in {"对外+对内", "对内+对外"}:
        return "对内/对外"
    return scope


def choose_source_docx() -> Path:
    candidates = [path for path in WORK_DIR.glob("*.docx") if not path.name.startswith("~$")]
    ranked: list[tuple[int, Path]] = []
    for path in candidates:
        try:
            doc = Document(path)
        except Exception:
            continue
        score = 0
        if len(doc.tables) > 9:
            table = doc.tables[9]
            table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
            if "报价明细表" in table_text:
                score += 10
            if "统一报价表" in table_text:
                score += 5
        if "替换稿" not in path.stem:
            score += 2
        ranked.append((score, path))
    ranked.sort(key=lambda item: item[0], reverse=True)
    if not ranked:
        raise FileNotFoundError("未找到可用的谈判文件模板")
    return ranked[0][1]


SOURCE_DOCX = choose_source_docx()


def set_run_font(run, size: float = 10.5) -> None:
    run.font.size = Pt(size)
    run.font.name = "宋体"
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def set_cell_text(cell, text: str, *, center: bool = False) -> None:
    cell.text = ""
    parts = text.split("\n") if text else [""]
    for idx, part in enumerate(parts):
        para = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        run = para.add_run(part)
        set_run_font(run)


def append_note(detail: str, note: str) -> str:
    note = normalize(note)
    if not note:
        return detail
    skip_tokens = ["暂不发布", "工作量", "人/天", "怎么计算", "怎么判定", "待提供", "无法按照"]
    if any(token in note for token in skip_tokens):
        return detail
    return f"{detail}\n补充说明：{note}"


def first_line(detail: str) -> str:
    lines = [line.strip() for line in detail.split("\n") if line.strip()]
    return lines[0] if lines else ""


def extract_named_requirement(detail: str) -> str:
    match = re.search(r"需求名称\s*\n([^\n]+)", detail)
    if match:
        return match.group(1).strip()
    return ""


def compact_module_title(title: str) -> str:
    title = normalize(title)
    title = re.sub(r"^\d+(?:\.\d+)*\s*", "", title).strip()
    title = title.lstrip(".").strip()
    return title


def pecms_platform_label(scope: str, module: str) -> str:
    scope = normalize_scope(scope)
    module = normalize(module)
    if module.startswith("开放接口方案"):
        return "对外PECMS开放接口开发"
    if scope == "对内":
        return "对内PECMS平台开发"
    if scope == "对外":
        return "对外PECMS平台开发"
    if scope == "对内/对外":
        return "对内/对外PECMS平台开发"
    return "PECMS平台开发"


def build_row_text(prefix: str, module: str, detail: str, note: str = "") -> str:
    module = normalize(module)
    detail = normalize(detail)
    if note:
        detail = append_note(detail, note)
    return f"{prefix}-{module}：{detail}"


def collect_pecms_rows(wb) -> list[str]:
    ws = wb.worksheets[1]
    rows: list[str] = []
    current_module = ""
    for excel_row in ws.iter_rows(min_row=6, values_only=True):
        module = normalize(excel_row[0])
        detail = normalize(excel_row[1])
        scope = normalize(excel_row[11])
        note = normalize(excel_row[13])
        if module in {"各项合计", "工作量汇总"}:
            break
        if module.startswith("附："):
            continue
        if module:
            current_module = module
        if not detail:
            continue
        prefix = pecms_platform_label(scope, current_module)
        rows.append(build_row_text(prefix, current_module, detail, note))
    return rows


def collect_app_rows(wb) -> list[str]:
    ws = wb.worksheets[2]
    rows: list[str] = []
    current_domain = ""
    current_module = ""
    prefix = "对内PECMS-APP端功能优化改造"
    for excel_row in ws.iter_rows(min_row=6, values_only=True):
        values = [normalize(item) for item in excel_row]
        if "各项合计" in values or "工作量汇总" in values:
            break
        if values[0]:
            current_domain = values[0]
        if values[1]:
            current_module = values[1]
        detail = values[2]
        if not detail:
            continue
        module = f"{current_domain}-{current_module}".strip("-") if current_domain else current_module
        rows.append(build_row_text(prefix, module, detail))
    return rows


def collect_safety_rows(wb) -> list[str]:
    ws = wb.worksheets[3]
    rows: list[str] = []
    current_module = ""
    prefix = "安全管理平台升级改造"
    for excel_row in ws.iter_rows(min_row=6, values_only=True):
        values = [normalize(item) for item in excel_row]
        if "各项合计" in values or "工作量汇总" in values:
            break
        if values[0]:
            current_module = values[0]
        detail = values[1]
        if not detail:
            continue
        named_requirement = extract_named_requirement(detail)
        if named_requirement and named_requirement != current_module:
            module = f"{current_module}-{named_requirement}" if current_module else named_requirement
        else:
            module = current_module
        note_parts = [values[7], values[9]]
        note = "\n".join(part for part in note_parts if part)
        rows.append(build_row_text(prefix, module, detail, note))
    return rows


def collect_om_rows(wb) -> list[str]:
    ws = wb.worksheets[4]
    rows: list[str] = []
    current_module = ""
    prefix = "执行管理与经营分析平台开发"
    for excel_row in ws.iter_rows(min_row=6, values_only=True):
        values = [normalize(item) for item in excel_row]
        if "各项合计" in values or "工作量汇总" in values:
            break
        if values[0]:
            current_module = values[0]
        detail = values[1]
        if not detail:
            continue
        sub_title = compact_module_title(first_line(detail))
        if sub_title and sub_title != current_module:
            module = f"{current_module}-{sub_title}"
        else:
            module = current_module
        rows.append(build_row_text(prefix, module, detail))
    return rows


def build_all_rows() -> list[str]:
    wb = load_workbook(SOURCE_XLSX, data_only=True)
    rows = []
    rows.extend(collect_pecms_rows(wb))
    rows.extend(collect_app_rows(wb))
    rows.extend(collect_safety_rows(wb))
    rows.extend(collect_om_rows(wb))
    return rows


def build_table_doc() -> Document:
    source = Document(SOURCE_DOCX)
    quote_table = source.tables[9]

    doc = Document()
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = Pt(10.5)

    cloned_tbl = copy.deepcopy(quote_table._tbl)
    doc._body._element.append(cloned_tbl)
    table = doc.tables[0]

    for _ in range(10):
        table._tbl.remove(table._tbl.tr_lst[0])

    return doc


def ensure_row_count(table, target_data_rows: int) -> None:
    current_data_rows = len(table.rows) - 2
    template_row = table.rows[-1]
    while current_data_rows < target_data_rows:
        table._tbl.append(copy.deepcopy(template_row._tr))
        current_data_rows += 1


def fill_table(doc: Document, row_texts: list[str]) -> None:
    table = doc.tables[0]
    ensure_row_count(table, len(row_texts))

    for row_index, text in enumerate(row_texts, start=2):
        row = table.rows[row_index]
        for col_index, cell in enumerate(row.cells):
            if col_index == 0:
                set_cell_text(cell, str(row_index - 1), center=True)
            elif col_index == 1:
                set_cell_text(cell, text)
            else:
                set_cell_text(cell, "")


def main() -> None:
    row_texts = build_all_rows()
    doc = build_table_doc()
    fill_table(doc, row_texts)
    doc.save(OUTPUT_DOCX)
    print(f"source_docx={SOURCE_DOCX}")
    print(f"source_xlsx={SOURCE_XLSX}")
    print(f"rows={len(row_texts)}")
    print(f"output={OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
