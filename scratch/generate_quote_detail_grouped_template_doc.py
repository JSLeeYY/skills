from __future__ import annotations

from collections import OrderedDict
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openpyxl import load_workbook


ROOT = Path.cwd()
DATE_DIR = ROOT / "word_project" / "2026" / "2026-04-24"
WORK_DIR = next(path for path in DATE_DIR.iterdir() if path.is_dir())
SOURCE_XLSX = next(path for path in WORK_DIR.glob("20260413*.xlsx"))
OUTPUT_DOCX = WORK_DIR / "分项报价表-按细分模块工作内容替换稿.docx"


def normalize(value: object) -> str:
    if value is None:
        return ""
    text = str(value).replace("\r\n", "\n").replace("\r", "\n").replace("\xa0", " ").strip()
    text = text.replace("\uf0b7", "- ")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def set_run_font(run, size: float = 10.5, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def set_paragraph_text(paragraph, text: str, *, align=WD_ALIGN_PARAGRAPH.LEFT, size: float = 10.5, bold: bool = False) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def set_cell_text(cell, text: str, *, align=WD_ALIGN_PARAGRAPH.LEFT, size: float = 10.5, bold: bool = False, vertical_center: bool = True) -> None:
    cell.text = ""
    parts = text.split("\n") if text else [""]
    for idx, part in enumerate(parts):
        paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        set_paragraph_text(paragraph, part, align=align, size=size, bold=bold)
    if vertical_center:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = Pt(10.5)


def build_sections() -> OrderedDict[str, OrderedDict[str, list[str]]]:
    wb = load_workbook(SOURCE_XLSX, data_only=True)
    sections: OrderedDict[str, OrderedDict[str, list[str]]] = OrderedDict()

    def add(section_title: str, module: str, detail: str) -> None:
        if not detail:
            return
        sections.setdefault(section_title, OrderedDict()).setdefault(module, []).append(detail)

    # PECMS: exclude open interface plans per latest user rule.
    pecms = wb.worksheets[1]
    current_module = ""
    for row in pecms.iter_rows(min_row=6, values_only=True):
        module = normalize(row[0])
        detail = normalize(row[1])
        scope = normalize(row[11])
        if module in {"各项合计", "工作量汇总"}:
            break
        if module.startswith("附：") or module.startswith("开放接口方案"):
            continue
        if module:
            current_module = module
        if not detail:
            continue
        if scope in {"对外+对内", "对内+对外"}:
            section_title = "PECMS模块（对内+对外）"
        elif scope == "对内":
            section_title = "PECMS模块（对内）"
        elif scope == "对外":
            section_title = "PECMS模块（对外）"
        else:
            section_title = "PECMS模块（通用）"
        add(section_title, current_module, detail)

    # APP
    app = wb.worksheets[2]
    current_domain = ""
    current_module = ""
    for row in app.iter_rows(min_row=6, values_only=True):
        values = [normalize(item) for item in row]
        if "各项合计" in values or "工作量汇总" in values:
            break
        if values[0]:
            current_domain = values[0]
        if values[1]:
            current_module = values[1]
        detail = values[2]
        if not detail:
            continue
        module = f"{current_domain}-{current_module}".strip("-")
        add("PECMS-APP端功能优化改造", module, detail)

    # Safety
    safety = wb.worksheets[3]
    current_module = ""
    for row in safety.iter_rows(min_row=6, values_only=True):
        values = [normalize(item) for item in row]
        if "各项合计" in values or "工作量汇总" in values:
            break
        if values[0]:
            current_module = values[0]
        detail = values[1]
        if not detail:
            continue
        add("安全管理平台", current_module, detail)

    # O&M
    om = wb.worksheets[4]
    current_module = ""
    for row in om.iter_rows(min_row=6, values_only=True):
        values = [normalize(item) for item in row]
        if "各项合计" in values or "工作量汇总" in values:
            break
        if values[0]:
            current_module = values[0]
        detail = values[1]
        if not detail:
            continue
        add("执行管理与经营分析平台", current_module, detail)

    ordered_titles = [
        "PECMS模块（对内+对外）",
        "PECMS模块（对内）",
        "PECMS模块（对外）",
        "PECMS模块（通用）",
        "PECMS-APP端功能优化改造",
        "安全管理平台",
        "执行管理与经营分析平台",
    ]
    ordered_sections: OrderedDict[str, OrderedDict[str, list[str]]] = OrderedDict()
    for title in ordered_titles:
        if title in sections:
            ordered_sections[title] = sections[title]
    return ordered_sections


def set_table_widths(table) -> None:
    widths = [Cm(1.6), Cm(4.2), Cm(10.4), Cm(3.2), Cm(3.0), Cm(2.8), Cm(1.8)]
    table.autofit = False
    for col_idx, width in enumerate(widths):
        for cell in table.columns[col_idx].cells:
            cell.width = width


def build_section_table(doc: Document, section_title: str, modules: OrderedDict[str, list[str]]) -> None:
    total_rows = 3 + sum(len(details) for details in modules.values())
    table = doc.add_table(rows=total_rows, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table)

    # Title row
    title_cell = table.rows[0].cells[0]
    for col_idx in range(1, 7):
        title_cell = title_cell.merge(table.rows[0].cells[col_idx])
    set_cell_text(table.rows[0].cells[0], "报价明细表", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=False)

    # Section row
    left_cell = table.rows[1].cells[0].merge(table.rows[1].cells[1])
    right_cell = table.rows[1].cells[2]
    for col_idx in range(3, 7):
        right_cell = right_cell.merge(table.rows[1].cells[col_idx])
    set_cell_text(left_cell, "工作量构成：", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    set_cell_text(right_cell, section_title, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

    # Header row
    headers = ["序\n号", "功能模块", "功能点描述", "人工日/数\n量", "单价（元）", "小计\n（元）", "备\n注"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[2].cells[idx], header, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

    # Data rows
    row_idx = 3
    serial = 1
    for module_name, details in modules.items():
        start = row_idx
        for detail in details:
            set_cell_text(table.rows[row_idx].cells[2], detail, align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5)
            set_cell_text(table.rows[row_idx].cells[3], "", align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(table.rows[row_idx].cells[4], "", align=WD_ALIGN_PARAGRAPH.CENTER)
            row_idx += 1
        end = row_idx - 1

        serial_cell = table.rows[start].cells[0]
        module_cell = table.rows[start].cells[1]
        subtotal_cell = table.rows[start].cells[5]
        remark_cell = table.rows[start].cells[6]
        if end > start:
            serial_cell = serial_cell.merge(table.rows[end].cells[0])
            module_cell = module_cell.merge(table.rows[end].cells[1])
            subtotal_cell = subtotal_cell.merge(table.rows[end].cells[5])
            remark_cell = remark_cell.merge(table.rows[end].cells[6])

        set_cell_text(serial_cell, str(serial), align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
        set_cell_text(module_cell, module_name, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
        set_cell_text(subtotal_cell, "", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(remark_cell, "", align=WD_ALIGN_PARAGRAPH.CENTER)
        serial += 1


def main() -> None:
    sections = build_sections()
    doc = Document()
    configure_document(doc)

    for idx, (section_title, modules) in enumerate(sections.items()):
        if idx > 0:
            doc.add_paragraph()
        build_section_table(doc, section_title, modules)

    doc.save(OUTPUT_DOCX)
    print(f"source_xlsx={SOURCE_XLSX}")
    print(f"tables={len(sections)}")
    print(f"output={OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
