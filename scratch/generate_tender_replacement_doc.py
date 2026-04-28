from __future__ import annotations

import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openpyxl import load_workbook


BASE_DIR = Path(
    r"D:\DevelopmentLocation\agent skill\skills\word_project\2026\2026-04-24\01-谈判文件对外发布版本内容编辑"
)
TENDER_DOC = BASE_DIR / "南京三方一体化数字信息平台优化提升项目谈判文件(2025-新增供应商考核)(修订稿).docx"
SCOREBOOK = BASE_DIR / "软件项目供应商评审评分表-最终版.xlsx"
OUTPUT_DOC = BASE_DIR / "南京三方一体化数字信息平台优化提升项目谈判文件第三章及第七章替换稿.docx"


def normalize_text(value: object) -> str:
    if value is None:
        return ""
    return str(value).replace("\r\n", "\n").replace("\xa0", " ").strip()


def set_run_style(run, *, size: float = 10.5, bold: bool = False):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.2)

    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Title"].font.size = Pt(18)
    doc.styles["Heading 1"].font.size = Pt(14)
    doc.styles["Heading 2"].font.size = Pt(12)
    doc.styles["Heading 3"].font.size = Pt(11)


def add_title(doc: Document, text: str, subtitle: str | None = None) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_run_style(run, size=18, bold=True)
    para.paragraph_format.space_after = Pt(8)

    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(subtitle)
        set_run_style(sub_run, size=10.5)
        sub_para.paragraph_format.space_after = Pt(8)


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    bold: bool = False,
    size: float = 10.5,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    set_run_style(run, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_heading(level=level)
    run = para.add_run(text)
    set_run_style(run, size=14 if level == 1 else 12 if level == 2 else 11, bold=True)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)


def shade_cell(cell, fill: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 9,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    set_run_style(run, size=size, bold=bold)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    column_widths: list[float] | None = None,
    center_cols: set[int] | None = None,
) -> None:
    center_cols = center_cols or set()
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(cell)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in center_cols else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], value, align=align)

    if column_widths:
        for row in table.rows:
            for idx, width in enumerate(column_widths):
                row.cells[idx].width = Cm(width)

    doc.add_paragraph("")


def add_signature_lines(doc: Document, items: list[str]) -> None:
    for item in items:
        add_paragraph(doc, item)


def extract_paragraphs(path: Path) -> list[str]:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    paragraphs: list[str] = []
    for p in root.findall(".//w:p", ns):
        texts = [node.text for node in p.findall(".//w:t", ns) if node.text]
        if texts:
            paragraphs.append("".join(texts).strip())
    return paragraphs


def locate_replace_ranges(paragraphs: list[str]) -> dict[str, int]:
    chapter3_start = paragraphs.index("第三章 谈判办法") + 1
    chapter4_start = paragraphs.index("第四章 技术规格及要求") + 1
    chapter7_positions = [idx + 1 for idx, text in enumerate(paragraphs) if text == "第七章 响应文件格式"]
    chapter8_start = paragraphs.index("第八章其他") + 1
    return {
        "chapter3_start": chapter3_start,
        "chapter3_end": chapter4_start - 1,
        "chapter7_cover_start": chapter7_positions[0],
        "chapter7_cover_end": chapter7_positions[1] - 1,
        "chapter7_replace_start": chapter7_positions[1],
        "chapter7_replace_end": chapter8_start - 1,
    }


def read_score_tables(path: Path) -> dict[str, list[list[str]]]:
    wb = load_workbook(path, data_only=False)

    def read_detail_sheet(sheet_name: str) -> list[list[str]]:
        ws = wb[sheet_name]
        rows: list[list[str]] = []
        current_major = ""
        for row_idx in range(5, ws.max_row + 1):
            major = normalize_text(ws[f"A{row_idx}"].value)
            sub_item = normalize_text(ws[f"B{row_idx}"].value)
            score = normalize_text(ws[f"C{row_idx}"].value)
            standard = normalize_text(ws[f"D{row_idx}"].value)
            evidence = normalize_text(ws[f"E{row_idx}"].value)
            if not any([major, sub_item, score, standard, evidence]):
                continue
            if "小计" in major or "合计" in major:
                break
            if major:
                current_major = major
            rows.append([current_major, sub_item, f"{score}分", standard, evidence])
        return rows

    weights = [
        ["1", "资质评分", "100分", "10%"],
        ["2", "定制开发评分", "100分", "40%"],
        ["3", "驻场运维评分", "100分", "20%"],
        ["4", "统一报价评分", "100分", "30%"],
        ["合计", "——", "——", "100%"],
    ]

    return {
        "weights": weights,
        "资质评分": read_detail_sheet("资质评分"),
        "定制开发评分": read_detail_sheet("定制开发评分"),
        "驻场运维评分": read_detail_sheet("驻场运维评分"),
    }


def build_replacement_scope(doc: Document, ranges: dict[str, int]) -> None:
    add_heading(doc, "一、替换范围及定位说明", level=1)
    add_paragraph(
        doc,
        "以下定位依据当前谈判文件正文段落抽取结果整理，用于明确原文替换范围。第三章建议整章整体替换；第七章建议保留封面页，替换清单及模板部分。",
    )
    add_table(
        doc,
        ["序号", "原文位置", "原文描述", "建议替换范围", "处理建议"],
        [
            [
                "1",
                f"第{ranges['chapter3_start']}-{ranges['chapter3_end']}段",
                "第三章《谈判办法》",
                "自“第三章 谈判办法”起至“3.6 谈判小组完成评审后……”止",
                "整章整体替换为与评分表完全对齐的综合评分法文本及明细表",
            ],
            [
                "2",
                f"第{ranges['chapter7_replace_start']}-{ranges['chapter7_replace_end']}段",
                "第七章《响应文件格式》中的清单及模板部分",
                "保留第七章封面页不动，自第二个“第七章 响应文件格式”标题起整体替换至“十三、供应商认为必要的其它文件”止",
                "重排响应文件清单，并补齐覆盖评分表全部评分点的模板和索引表",
            ],
        ],
        column_widths=[1.2, 2.8, 3.1, 6.4, 4.8],
        center_cols={0},
    )


def build_gap_analysis(doc: Document) -> None:
    add_heading(doc, "二、评分表对照与缺项结论", level=1)
    add_paragraph(
        doc,
        "对照《软件项目供应商评审评分表-最终版.xlsx》后，现有谈判文件主要存在两类问题：一是第三章只写到了评分大项和权重，没有把评分表中的分档标准、证明材料和评分依据完整落到文本；二是第七章虽然补入了部分模板，但仍缺少足以支撑全部评分点取证和检索的目录、索引和专项承诺模板。",
    )
    comparison_rows = [
        [
            "资质评分：企业规模、财务状况、纳税能力",
            "现有第七章仅笼统要求提供财务报表、纳税证明，没有结构化模板，容易出现材料零散、专家难以快速定位的问题。",
            "新增“供应商基本情况表”“近三年财务状况表”“纳税及开票能力情况说明表”，并在资格目录中要求标注页码。",
        ],
        [
            "资质评分：信用合规、材料真实性",
            "现有文本缺少“材料真实性及履约承诺”模板，信用合规材料与证明页码也未建立索引。",
            "新增“信用合规承诺函”“材料真实性及履约承诺函”，并要求目录索引对应页码。",
        ],
        [
            "资质评分：基础资质、安全保密、专门条件响应",
            "现有文本虽列举部分资质，但没有形成专项清单，不利于审查是否齐全。",
            "在“资格及资质证明文件目录及索引表”中单列专项资质、保密、安全、授权、研发能力等证明材料。",
        ],
        [
            "定制开发评分：需求理解、架构、接口、数据、安全",
            "现有第三章只列分值和大项名称，缺少分档标准；第七章也缺少技术方案响应索引。",
            "第三章按评分表逐项补入分档标准和证明材料；第七章新增“技术方案响应索引表”，要求逐项对应页码。",
        ],
        [
            "定制开发评分：实施计划、风险、质量、交付物、验收安排",
            "现有模板未把阶段交付、验收资料和评分点建立清晰映射。",
            "新增“项目里程碑计划表”“交付物清单”“验收配合安排表”“风险识别及应对措施表”。",
        ],
        [
            "定制开发评分：项目团队配置与履约能力",
            "现有团队模板较粗，缺少职责、证明材料页码和驻场周期等信息。",
            "补齐“项目团队人员配置表”“核心人员简历表”“驻场安排及稳定性保障表”，并要求附岗位证明材料页码。",
        ],
        [
            "驻场运维评分：服务流程、SLA、应急处置、服务治理",
            "现有模板只覆盖部分承诺，未形成完整的服务治理、例会、报告和升级闭环模板。",
            "新增“服务响应及应急处置承诺表”“服务治理与持续优化机制表”“运维知识库及问题闭环管理表”。",
        ],
        [
            "驻场运维评分：商务响应与文件完整性",
            "现有文本缺少总目录、版本记录和商务响应矩阵，不利于支撑“文件完整性、结构清晰、检索方便”评分。",
            "新增“响应文件目录及页码索引表”“文件版本记录表”“商务响应汇总表”“采购单位关注事项响应表”。",
        ],
    ]
    add_table(
        doc,
        ["评分表关注点", "现有文档问题", "本次补齐方式"],
        comparison_rows,
        column_widths=[4.4, 6.5, 7.1],
    )


def build_chapter3(doc: Document, score_tables: dict[str, list[list[str]]], ranges: dict[str, int]) -> None:
    add_heading(doc, "三、建议替换的第三章文本", level=1)
    add_paragraph(
        doc,
        f"建议将原谈判文件第{ranges['chapter3_start']}-{ranges['chapter3_end']}段整体替换为以下内容。下文已按评分表将评分标准、分档描述、证明材料、价格公式和排名规则完整展开，可直接作为第三章正文使用。",
    )
    add_heading(doc, "第三章 谈判办法", level=1)

    add_paragraph(doc, "（一）谈判方法", bold=True)
    add_paragraph(
        doc,
        "1. 本次谈判采用综合评分法确定候选成交供应商。谈判小组对通过初步评审的供应商，按照本章规定的评审因素、评分标准及权重进行综合评分，并按照综合评审得分由高到低的顺序推荐候选成交供应商。",
    )
    add_paragraph(
        doc,
        "2. 本次谈判不以最低报价作为唯一成交依据。供应商报价作为综合评分的重要组成部分，与供应商资质能力、定制开发方案、驻场运维保障能力、商务响应情况等因素共同参与综合评审。",
    )
    add_paragraph(
        doc,
        "3. 谈判小组可以根据项目实际需要，对供应商的响应文件、技术方案、人员配置、驻场安排、服务承诺、报价构成及其他相关事项进行谈判、澄清或要求供应商作出书面承诺。供应商的澄清、说明和承诺不得超出响应文件的范围或者改变响应文件的实质性内容；经谈判小组确认的书面澄清、说明和承诺，作为响应文件的组成部分。",
    )

    add_paragraph(doc, "（二）评审标准", bold=True)
    add_paragraph(doc, "1. 初步评审", bold=True)
    add_paragraph(doc, "1.1 初步评审程序", bold=True)
    add_paragraph(
        doc,
        "谈判小组对供应商提交的响应文件进行初步评审。初步评审包括资格性审查、符合性审查、响应文件完整性审查、报价有效性审查及谈判文件要求的其他基础审查事项。初步评审中有一项未通过的，否决其响应文件，该供应商不得进入详细评分和综合排名。",
    )
    add_paragraph(doc, "1.2 报价修正", bold=True)
    add_paragraph(
        doc,
        "响应文件的报价有算术错误的，谈判小组按以下原则对报价进行修正，修正后的价格经供应商书面确认后具有约束力。供应商不接受修正价格的，其响应文件按无效处理：",
    )
    add_paragraph(doc, "（1）报价大写金额与小写金额不一致的，以大写金额为准；")
    add_paragraph(doc, "（2）报价总价金额与单价金额计算结果不一致的，以单价金额为准并修正总价；单价金额小数点有明显错误的除外；")
    add_paragraph(doc, "（3）分项报价汇总金额与报价总价不一致的，以分项报价汇总金额为基础修正报价总价，但谈判小组认为存在明显计算或填写错误的除外。")
    add_paragraph(doc, "1.3 初步评审标准", bold=True)
    add_table(
        doc,
        ["序号", "评审因素", "评审标准"],
        [
            ["1", "营业执照", "供应商名称应与营业执照、响应文件及相关证明材料一致；供应商应为在中华人民共和国境内注册、具有独立承担民事责任能力的法人或其他组织，经营范围或业务能力应符合本项目采购需求。"],
            ["2", "签字盖章", "响应文件中需加盖公章、法定代表人签章或委托代理人签字处，应严格按照谈判文件规定执行。"],
            ["3", "资质文件", "供应商应具备软件开发、测试、实施、运维及项目管理所需的相应能力，并按谈判文件要求提供相关证明材料。供应商应具备一定数量从事软件开发、测试、实施、运维和项目管理的专业技术人员。"],
            ["4", "业绩要求", "供应商应具有软件开发、系统建设、平台优化、驻场运维或相关服务业绩。如供应商承接过类似项目，熟悉相关业务流程并能够提供有效证明材料的，作为后续详细评分依据。"],
            ["5", "商务及服务承诺", "供应商应对人员配置、入场时间、驻场安排、开发维护参与度、需求沟通响应、技术方案、付款条件、保密要求、安全要求、知识产权、违约责任等谈判文件要求作出实质性响应。"],
            ["6", "联合体要求", "本次谈判不接受供应商联合体参与报价。"],
        ],
        column_widths=[1.2, 3.0, 12.8],
        center_cols={0},
    )

    add_paragraph(doc, "2. 详细评分", bold=True)
    add_paragraph(doc, "2.1 评分组成及权重", bold=True)
    add_paragraph(
        doc,
        "对通过初步评审的供应商，谈判小组按照资质评分、定制开发评分、驻场运维评分、统一报价评分四个部分进行评分。各部分均按100分制独立评分，再按权重折算为综合评审得分。",
    )
    add_paragraph(doc, "综合评审得分=资质评分×10%+定制开发评分×40%+驻场运维评分×20%+统一报价评分×30%。")
    add_table(
        doc,
        ["序号", "评分项目", "满分", "权重"],
        score_tables["weights"],
        column_widths=[1.4, 5.3, 3.0, 2.0],
        center_cols={0, 2, 3},
    )

    add_paragraph(doc, "2.2 资质评分", bold=True)
    add_paragraph(
        doc,
        "资质评分满分100分，权重10%。资质评分按下表执行；供应商提供的证明材料应与响应文件目录及页码索引对应，未提供有效证明材料、证明材料不能支撑评分项，或者材料与响应内容明显不一致的，相应评分项可酌情扣分或不得分。",
    )
    add_table(
        doc,
        ["评审大项", "评审子项", "分值", "评分标准", "评分依据/证明材料"],
        score_tables["资质评分"],
        column_widths=[3.4, 3.8, 1.6, 7.8, 4.8],
    )

    add_paragraph(doc, "2.3 定制开发评分", bold=True)
    add_paragraph(
        doc,
        "定制开发评分满分100分，权重40%。定制开发评分按下表执行；技术方案、实施计划、交付安排和团队配置应与本项目实际范围逐项对应，未提供有效证明材料、方案明显不符合项目需求的，相应评分项可酌情扣分或不得分。",
    )
    add_table(
        doc,
        ["评审大项", "评审子项", "分值", "评分标准", "评分依据/证明材料"],
        score_tables["定制开发评分"],
        column_widths=[3.4, 3.8, 1.6, 7.8, 4.8],
    )

    add_paragraph(doc, "2.4 驻场运维评分", bold=True)
    add_paragraph(
        doc,
        "驻场运维评分满分100分，权重20%。驻场运维评分按下表执行；服务方案、驻场安排、SLA 响应、知识转移和服务治理等内容应形成明确承诺并提供对应证明材料，未提供有效证明材料或承诺明显不能满足项目需要的，相应评分项可酌情扣分或不得分。",
    )
    add_table(
        doc,
        ["评审大项", "评审子项", "分值", "评分标准", "评分依据/证明材料"],
        score_tables["驻场运维评分"],
        column_widths=[3.4, 3.8, 1.6, 7.8, 4.8],
    )

    add_paragraph(doc, "2.5 统一报价评分", bold=True)
    add_paragraph(
        doc,
        "统一报价评分满分100分，权重30%。供应商应按本项目完整范围提交一份整包统一有效报价，不再区分定制开发和驻场运维分别报价。价格评分采用最低有效报价比例法。满足谈判文件要求且价格有效的供应商中，最低有效报价得100分；其他供应商报价得分按以下公式计算：",
    )
    add_paragraph(doc, "统一报价得分=最低有效报价/本供应商有效报价×100。")
    add_paragraph(
        doc,
        "计算结果保留两位小数。供应商报价明显低于其他有效供应商报价，或者明显低于采购成本、可能影响项目质量和诚信履约的，谈判小组有权要求供应商在规定时间内作出书面说明并提供相关证明材料。供应商不能合理说明报价合理性，或者不能证明其能够按照谈判文件要求履约的，其报价可被认定为异常低价，谈判小组可否决其响应文件或在评审中不予采信。",
    )
    add_table(
        doc,
        ["评分项目", "分值", "评分标准", "评分依据/证明材料"],
        [
            [
                "统一报价得分",
                "100分",
                "价格部分采用最低有效报价比例法，最低有效报价得满分100分，其他供应商按“最低有效报价/本供应商有效报价×100”折算。",
                "有效报价表、报价函、统一报价汇总表、分项报价明细表、澄清函及其他价格说明材料。",
            ]
        ],
        column_widths=[3.4, 1.8, 8.7, 6.1],
        center_cols={1},
    )

    add_paragraph(doc, "3. 评分及排名规则", bold=True)
    add_paragraph(doc, "3.1 谈判小组成员应依据谈判文件、响应文件、供应商澄清说明、承诺文件及有效证明材料独立评分。")
    add_paragraph(doc, "3.2 各评分项得分不得超过该项满分。供应商未提供相应证明材料、证明材料不完整、证明材料不能支撑其响应内容，或者响应内容与本项目需求明显不匹配的，谈判小组可根据评分标准酌情扣分或不得分。")
    add_paragraph(doc, "3.3 综合评审得分按以下公式计算：综合评审得分=资质评分×10%+定制开发评分×40%+驻场运维评分×20%+统一报价评分×30%。")
    add_paragraph(doc, "3.4 综合评审得分计算结果保留两位小数。谈判小组按综合评审得分由高到低进行排序，并推荐候选成交供应商。")
    add_paragraph(doc, "3.5 综合评审得分相同的，按统一报价得分高者优先；统一报价得分仍相同的，按定制开发评分高者优先；定制开发评分仍相同的，按驻场运维评分高者优先；仍无法确定排序的，由谈判小组结合供应商履约能力、服务承诺、项目风险等因素集体讨论确定推荐顺序。")
    add_paragraph(doc, "3.6 谈判小组完成评审后，应形成评审意见，列明通过初步评审的供应商、未通过初步评审的供应商及原因、各供应商综合得分、排名及候选成交供应商推荐意见。承办部门根据谈判小组推荐结果，报请公司批准后确定成交供应商。")


def build_chapter7(doc: Document, ranges: dict[str, int]) -> None:
    add_heading(doc, "四、建议替换的第七章文本", level=1)
    add_paragraph(
        doc,
        f"建议保留原第七章封面页（第{ranges['chapter7_cover_start']}-{ranges['chapter7_cover_end']}段）不动，自第{ranges['chapter7_replace_start']}段“第七章 响应文件格式”起整体替换至第{ranges['chapter7_replace_end']}段结束。以下内容已将评分表要求映射为清单、索引和模板，可直接作为第七章正文使用。",
    )
    add_heading(doc, "第七章 响应文件格式", level=1)

    add_paragraph(doc, "响应文件编制说明", bold=True)
    add_paragraph(doc, "1. 响应文件应按照本章规定的顺序编制，并设置连续页码。目录、索引表、正文、附件和证明材料应对应一致，便于谈判小组检索。")
    add_paragraph(doc, "2. 对于同一份证明材料支撑多个评分项的，供应商可以重复引用同一材料，但应在相应目录、索引表或说明表中明确对应页码。")
    add_paragraph(doc, "3. 供应商对评分项仅作承诺但未提供相应证明材料、样稿、制度、案例或其他支撑内容的，相关评分项可以扣分或不得分。")
    add_paragraph(doc, "4. 供应商可以在不减少本章要求内容的前提下，对模板进行扩展、补充或优化，但不得删减实质性响应内容。")

    add_paragraph(doc, "文件格式按下列顺序编制（清单）", bold=True)
    checklist_rows = [
        ["1", "响应文件目录及页码索引表、文件版本记录表", "是", "文件完整性、结构清晰、检索方便", "总目录、分册页码、版本号、修订说明"],
        ["2", "报价函", "是", "初步评审、统一报价评分", "统一总报价、交付期、有效期、总体承诺"],
        ["3", "统一报价汇总表", "是", "统一报价评分", "整包统一报价、费用构成说明"],
        ["4", "分项报价明细表及价格构成说明", "是", "统一报价评分、报价合理性说明", "分项构成、工作量、单价、小计、说明"],
        ["5", "资格及资质证明文件目录及索引表", "是", "资质评分、初步评审", "资格资质、页码、备注"],
        ["6", "供应商基本情况表", "是", "企业规模与综合实力", "成立时间、注册资本、人员规模、主营业务、核心优势"],
        ["7", "近三年财务状况表及纳税开票能力说明", "是", "财务状况、企业规模与综合实力", "收入、利润、资产、负债、纳税及开票能力"],
        ["8", "信用合规承诺函、材料真实性及履约承诺函", "是", "企业信誉与合规情况、文件完整性", "无重大违法、无重大质量安全事故、材料真实性承诺"],
        ["9", "供应商代表身份证明", "是", "初步评审", "法定代表人身份证明"],
        ["10", "授权委托书（如适用）", "条件提供", "初步评审", "授权范围、期限、被授权人信息"],
        ["11", "同类项目业绩证明材料", "是", "同类项目业绩、同类软件交付经验、同类项目案例", "业绩汇总表、单项业绩说明表、合同关键页、验收证明等"],
        ["12", "技术方案及响应索引表", "是", "需求理解、边界识别、架构、接口、数据、安全、可扩展性", "章节说明、索引页码、图表样稿"],
        ["13", "项目实施计划及交付方案", "是", "实施计划、风险与质量控制、交付物定义与验收安排", "里程碑计划、交付物清单、验收配合、风险表"],
        ["14", "项目团队及驻场服务方案", "是", "项目团队配置与履约能力、驻场人员配置与稳定性保障", "团队表、简历表、驻场安排、替补机制"],
        ["15", "运维支持、响应保障及知识转移方案", "是", "服务理解、服务可执行性、响应时效、知识交接、服务治理", "SLA、应急表、治理机制、知识库、培训移交计划"],
        ["16", "商务响应及专项承诺文件", "是", "基础资质响应、商务响应与文件完整性", "采购单位关注事项响应表、商务响应汇总表、服务承诺、保密及知识产权承诺"],
        ["17", "商务条款偏离表", "是", "商务响应与文件完整性", "逐条对比商务条款是否偏离"],
        ["18", "技术规格偏离表", "是", "技术响应、文件完整性", "逐条对比技术要求是否偏离"],
        ["19", "供应商认为必要的其他文件", "可选", "补充说明", "其他有利于评审的证明材料"],
    ]
    add_table(
        doc,
        ["序号", "文件组成", "是否必需", "对应评分/审查项", "主要内容或附件"],
        checklist_rows,
        column_widths=[1.1, 5.2, 1.8, 4.6, 6.2],
        center_cols={0, 2},
    )

    add_paragraph(doc, "一、响应文件目录及页码索引表", bold=True)
    add_table(
        doc,
        ["序号", "文件名称", "页码范围", "对应评分项/审查项", "备注"],
        [["1", "", "", "", ""], ["2", "", "", "", ""], ["3", "", "", "", ""], ["4", "", "", "", ""], ["5", "", "", "", ""]],
        column_widths=[1.1, 5.4, 2.5, 6.3, 3.0],
        center_cols={0},
    )
    add_paragraph(doc, "文件版本记录表", bold=True)
    add_table(
        doc,
        ["版本号", "编制日期", "编制/修订人员", "修订内容说明"],
        [["V1.0", "", "", ""], ["V1.1", "", "", ""], ["V1.2", "", "", ""]],
        column_widths=[2.0, 3.0, 4.0, 9.3],
    )

    add_paragraph(doc, "二、报价函", bold=True)
    add_paragraph(doc, "致：南京三方化工设备监理有限公司")
    add_paragraph(
        doc,
        "根据已收到的《南京三方一体化数字信息平台优化提升项目》谈判文件，我方经认真研究后，决定参加本项目谈判，并授权________（姓名、职务）代表我方提交响应文件。在此，我方郑重承诺如下：",
    )
    add_paragraph(doc, "1. 我方愿意按谈判文件、响应文件、澄清承诺及合同约定承担本项目全部工作内容，并提交针对本项目完整范围的整包统一报价。")
    add_paragraph(doc, "2. 我方统一总报价为：人民币（大写）____________________；人民币（小写）￥____________________。")
    add_paragraph(doc, "3. 上述报价已覆盖本项目全部开发、实施、驻场、运维、培训、资料移交、质保及谈判文件要求的其他服务内容。")
    add_paragraph(doc, "4. 我方承诺交付期、驻场安排、服务响应和验收配合均满足谈判文件要求，并按约履行合同义务。")
    add_paragraph(doc, "5. 我方已详细审查全部谈判文件，完全理解并接受谈判文件及合同条款的全部要求。")
    add_paragraph(doc, "6. 本响应文件自谈判之日起60日历天内有效。")
    add_paragraph(doc, "7. 我方同意按采购单位要求提供与本次谈判有关的全部资料、说明和澄清文件。")
    add_signature_lines(
        doc,
        [
            "供应商名称（盖公章）：____________________",
            "法定代表人或委托代理人（签字）：____________________",
            "联系人：____________________    联系电话：____________________",
            "日期：______年______月______日",
        ],
    )

    add_paragraph(doc, "三、统一报价汇总表", bold=True)
    add_paragraph(doc, "供应商应提交针对本项目完整范围的整包统一报价。下表为建议填报格式，供应商可根据实际情况扩展行次，但不得拆分为多份独立总报价。")
    add_table(
        doc,
        ["序号", "费用项", "报价（元）", "说明"],
        [
            ["1", "经营管理分析模块开发", "", ""],
            ["2", "PECMS（对内+对外PC）模块优化改造", "", ""],
            ["3", "APP端功能优化改造", "", ""],
            ["4", "安全管理模块优化改造", "", ""],
            ["5", "驻场服务与运维保障", "", ""],
            ["6", "培训、知识转移及资料移交", "", ""],
            ["7", "质保及售后服务", "", ""],
            ["8", "其他", "", ""],
            ["合计", "统一总报价", "", ""],
        ],
        column_widths=[1.2, 7.2, 3.0, 6.4],
        center_cols={0, 2},
    )

    add_paragraph(doc, "四、分项报价明细表及价格构成说明", bold=True)
    add_paragraph(doc, "供应商应对统一总报价的构成进行必要分解说明，用于评审报价合理性，但最终评审仍以统一总报价为准。")
    add_table(
        doc,
        ["序号", "工作内容/费用项", "工作量说明", "人工日/数量", "单价（元）", "小计（元）", "备注"],
        [["1", "", "", "", "", "", ""], ["2", "", "", "", "", "", ""], ["3", "", "", "", "", "", ""], ["4", "", "", "", "", "", ""], ["5", "", "", "", "", "", ""]],
        column_widths=[1.1, 4.8, 4.6, 2.3, 2.6, 2.6, 3.1],
        center_cols={0, 3, 4, 5},
    )
    add_paragraph(doc, "统一报价与分项报价一致性说明：______________________________________________________________")

    add_paragraph(doc, "五、资格及资质证明文件", bold=True)
    add_paragraph(doc, "资格及资质证明文件目录及索引表", bold=True)
    add_table(
        doc,
        ["序号", "材料名称", "是否提供", "页码", "备注"],
        [
            ["1", "营业执照", "□是  □否", "", ""],
            ["2", "开票能力证明", "□是  □否", "", ""],
            ["3", "企业基本情况说明", "□是  □否", "", ""],
            ["4", "近三年审计报告或财务报表", "□是  □否", "", ""],
            ["5", "纳税证明", "□是  □否", "", ""],
            ["6", "信用承诺或信用查询结果", "□是  □否", "", ""],
            ["7", "无重大违法记录声明", "□是  □否", "", ""],
            ["8", "无重大质量安全事故声明", "□是  □否", "", ""],
            ["9", "保密承诺及安全管理制度", "□是  □否", "", ""],
            ["10", "专项资质、授权文件", "□是  □否", "", ""],
            ["11", "软件著作权、专利或研发能力证明", "□是  □否", "", ""],
            ["12", "服务管理制度、SLA治理机制说明", "□是  □否", "", ""],
            ["13", "其他与资质评分有关的证明材料", "□是  □否", "", ""],
        ],
        column_widths=[1.1, 7.0, 2.4, 2.2, 4.7],
        center_cols={0, 2, 3},
    )

    add_paragraph(doc, "供应商基本情况表", bold=True)
    add_table(
        doc,
        ["项目", "填写内容"],
        [
            ["供应商名称", ""],
            ["成立时间", ""],
            ["注册资本", ""],
            ["企业性质", ""],
            ["注册地址及主要办公地点", ""],
            ["员工总人数及研发/实施/运维人员规模", ""],
            ["主营业务及与本项目相关的核心能力", ""],
            ["近三年主要经营情况概述", ""],
            ["与本项目相关的典型优势", ""],
            ["联系人及联系方式", ""],
        ],
        column_widths=[4.2, 13.2],
    )

    add_paragraph(doc, "近三年财务状况表（建议按2023-2025年度填写）", bold=True)
    add_table(
        doc,
        ["项目", "2023年度", "2024年度", "2025年度", "备注"],
        [
            ["主营业务收入（元）", "", "", "", ""],
            ["净利润（元）", "", "", "", ""],
            ["资产总额（元）", "", "", "", ""],
            ["负债总额（元）", "", "", "", ""],
            ["资产负债率（%）", "", "", "", ""],
            ["经营现金流情况", "", "", "", ""],
            ["其他需说明事项", "", "", "", ""],
        ],
        column_widths=[4.6, 3.1, 3.1, 3.1, 3.5],
    )

    add_paragraph(doc, "纳税及开票能力情况说明表", bold=True)
    add_table(
        doc,
        ["项目", "填写内容"],
        [
            ["纳税主体名称", ""],
            ["纳税人类型", ""],
            ["可开具发票类型及税率", ""],
            ["近三年主要纳税情况说明", ""],
            ["与本项目相关的开票安排说明", ""],
            ["相关证明材料页码", ""],
        ],
        column_widths=[5.2, 12.2],
    )

    add_paragraph(doc, "信用合规承诺函", bold=True)
    add_paragraph(doc, "致：南京三方化工设备监理有限公司")
    add_paragraph(
        doc,
        "我公司郑重承诺：近三年内不存在重大违法失信记录，不存在重大质量安全事故，不存在严重违约情形；本次提交的信用、合规及资质材料真实、完整、可核验。如有不实，我公司愿承担由此产生的一切责任。",
    )
    add_signature_lines(
        doc,
        [
            "供应商名称（盖公章）：____________________",
            "法定代表人或委托代理人（签字）：____________________",
            "日期：______年______月______日",
        ],
    )

    add_paragraph(doc, "材料真实性及履约承诺函", bold=True)
    add_paragraph(doc, "致：南京三方化工设备监理有限公司")
    add_paragraph(
        doc,
        "我公司承诺本响应文件及所附全部证明材料真实、准确、完整，不存在伪造、变造、隐瞒或重大遗漏。若成交，我公司将严格按照谈判文件、响应文件、澄清承诺和合同约定完成项目开发、驻场、运维、培训、资料移交及质保工作。如有虚假材料或无法履约情形，我公司愿承担相应法律责任和违约责任。",
    )
    add_signature_lines(
        doc,
        [
            "供应商名称（盖公章）：____________________",
            "法定代表人或委托代理人（签字）：____________________",
            "日期：______年______月______日",
        ],
    )

    add_paragraph(doc, "六、供应商代表身份证明", bold=True)
    add_paragraph(doc, "法定代表人身份证明（法定代表人参加谈判时适用）", bold=True)
    add_table(
        doc,
        ["项目", "填写内容"],
        [
            ["供应商名称", ""],
            ["单位性质", ""],
            ["地址", ""],
            ["成立时间", ""],
            ["经营期限", ""],
            ["法定代表人姓名", ""],
            ["性别/年龄/职务", ""],
            ["身份证号码", ""],
        ],
        column_widths=[4.6, 13.0],
    )
    add_signature_lines(
        doc,
        [
            "附：法定代表人身份证复印件（正反面）",
            "供应商名称（盖公章）：____________________",
            "日期：______年______月______日",
        ],
    )

    add_paragraph(doc, "七、授权委托书（授权委托人参加谈判时适用）", bold=True)
    add_paragraph(
        doc,
        "本人________系________（供应商名称）的法定代表人，现委托________（姓名、职务）为我方代理人。代理人根据授权，以我方名义签署、澄清、说明、补正、递交、撤回、修改本项目响应文件和处理有关事宜，其法律后果由我方承担。",
    )
    add_paragraph(doc, "委托期限：自本授权委托书签署之日起至本次谈判结束并完成相关后续事宜止。")
    add_paragraph(doc, "代理人无转委托权。")
    add_signature_lines(
        doc,
        [
            "供应商名称（盖公章）：____________________",
            "法定代表人（签字或盖章）：____________________",
            "委托代理人（签字）：____________________",
            "附：法定代表人、委托代理人身份证复印件（正反面）",
            "日期：______年______月______日",
        ],
    )

    add_paragraph(doc, "八、同类项目业绩证明材料", bold=True)
    add_paragraph(doc, "同类项目业绩汇总表", bold=True)
    add_table(
        doc,
        ["序号", "项目名称", "委托单位", "项目类型", "合同金额", "实施/服务周期", "主要工作内容", "验收/履约情况", "证明材料页码"],
        [["1", "", "", "", "", "", "", "", ""], ["2", "", "", "", "", "", "", "", ""], ["3", "", "", "", "", "", "", "", ""]],
        column_widths=[1.0, 3.0, 2.9, 2.1, 2.2, 2.3, 3.8, 2.5, 2.0],
        center_cols={0},
    )
    add_paragraph(doc, "单项业绩说明表（每项业绩可复制填写）", bold=True)
    add_table(
        doc,
        ["项目", "填写内容"],
        [
            ["项目名称", ""],
            ["委托单位", ""],
            ["项目合同金额", ""],
            ["项目实施周期", ""],
            ["项目类型", "□软件开发  □系统建设  □平台优化  □驻场运维  □其他："],
            ["项目背景及业务场景", ""],
            ["供应商承担的主要工作", ""],
            ["与本项目的相似性说明", ""],
            ["交付成果", ""],
            ["验收或履约情况", ""],
            ["所附证明材料清单", "□合同关键页  □验收证明  □用户评价  □项目成果说明  □发票/回款证明  □其他："],
        ],
        column_widths=[4.7, 12.9],
    )

    add_paragraph(doc, "九、技术方案", bold=True)
    add_paragraph(doc, "供应商应结合本项目需求编制技术方案。技术方案不限定页数，但应至少覆盖以下内容，并与谈判文件第四章技术规格及要求逐项呼应。")
    add_table(
        doc,
        ["序号", "章节名称", "应说明的主要内容", "对应评分关注点"],
        [
            ["1", "项目理解与需求分析", "项目背景、建设目标、业务场景、核心需求、现状问题、目标拆解。", "需求理解与场景匹配度"],
            ["2", "范围边界与关键难点识别", "功能边界、接口边界、数据边界、实施边界、关键难点及应对思路。", "范围边界与关键难点识别"],
            ["3", "总体架构与功能方案", "总体架构、功能模块、业务流程、系统交互、部署架构、技术路线。", "总体架构与功能方案合理性"],
            ["4", "接口、数据与安全设计", "接口清单、数据流向、数据模型、权限控制、安全设计、日志审计。", "接口、数据与安全设计"],
            ["5", "可扩展性与兼容性设计", "可扩展方案、兼容性方案、现有系统适配、移动端/PC端适配。", "可扩展性与可实施性"],
            ["6", "测试与质量保障方案", "功能测试、性能测试、安全测试、兼容性测试、缺陷管理、质量控制。", "实施计划、风险与质量控制"],
            ["7", "交付成果说明", "源码、数据库脚本、接口文档、部署说明、用户手册、运维手册、测试报告等。", "交付物定义与验收安排"],
        ],
        column_widths=[1.0, 3.4, 9.2, 5.0],
        center_cols={0},
    )
    add_paragraph(doc, "技术方案响应索引表", bold=True)
    add_table(
        doc,
        ["评分项/需求项", "响应文件章节", "页码", "备注"],
        [
            ["项目理解与需求分析", "", "", ""],
            ["范围边界与关键难点识别", "", "", ""],
            ["总体架构与功能方案", "", "", ""],
            ["接口、数据与安全设计", "", "", ""],
            ["可扩展性与兼容性设计", "", "", ""],
            ["测试与质量保障方案", "", "", ""],
            ["交付成果说明", "", "", ""],
        ],
        column_widths=[5.0, 4.8, 2.2, 6.0],
    )

    add_paragraph(doc, "十、项目实施计划及交付方案", bold=True)
    add_paragraph(doc, "项目里程碑计划表", bold=True)
    add_table(
        doc,
        ["序号", "阶段", "计划起止时间", "主要工作内容", "阶段交付成果", "甲方配合事项", "验收/确认方式"],
        [
            ["1", "需求澄清与计划确认", "", "", "", "", ""],
            ["2", "第一阶段开发与交付", "", "", "", "", ""],
            ["3", "第二阶段开发与交付", "", "", "", "", ""],
            ["4", "第三阶段开发与交付", "", "", "", "", ""],
            ["5", "联调测试与问题整改", "", "", "", "", ""],
            ["6", "终验及资料移交", "", "", "", "", ""],
        ],
        column_widths=[1.0, 3.0, 2.6, 4.4, 4.0, 3.0, 2.7],
        center_cols={0},
    )
    add_paragraph(doc, "交付物清单", bold=True)
    add_table(
        doc,
        ["序号", "交付物名称", "交付形式", "对应阶段", "验收标准/确认方式", "备注"],
        [
            ["1", "源代码（含注释）", "", "", "", ""],
            ["2", "数据库脚本", "", "", "", ""],
            ["3", "接口文档", "", "", "", ""],
            ["4", "部署说明", "", "", "", ""],
            ["5", "系统运维手册", "", "", "", ""],
            ["6", "用户操作手册", "", "", "", ""],
            ["7", "测试报告", "", "", "", ""],
            ["8", "培训资料", "", "", "", ""],
            ["9", "缺陷清单及整改闭环记录", "", "", "", ""],
        ],
        column_widths=[1.0, 4.5, 2.4, 2.6, 6.0, 2.4],
        center_cols={0},
    )
    add_paragraph(doc, "验收配合安排表", bold=True)
    add_table(
        doc,
        ["阶段", "拟提交验收资料", "供应商自检要求", "甲方确认方式", "计划时间"],
        [["需求澄清", "", "", "", ""], ["阶段一", "", "", "", ""], ["阶段二", "", "", "", ""], ["阶段三", "", "", "", ""], ["终验", "", "", "", ""]],
        column_widths=[2.4, 5.8, 4.6, 3.4, 2.4],
    )
    add_paragraph(doc, "风险识别及应对措施表", bold=True)
    add_table(
        doc,
        ["序号", "风险事项", "可能影响", "预防措施", "应急处置措施", "责任岗位"],
        [["1", "", "", "", "", ""], ["2", "", "", "", "", ""], ["3", "", "", "", "", ""]],
        column_widths=[1.0, 4.0, 3.4, 4.2, 4.2, 2.0],
        center_cols={0},
    )

    add_paragraph(doc, "十一、项目团队及驻场服务方案", bold=True)
    add_paragraph(doc, "项目团队人员配置表", bold=True)
    add_table(
        doc,
        ["序号", "姓名", "岗位/角色", "拟投入工作内容", "驻场/远程", "计划投入时间", "相关经验", "证明材料页码"],
        [
            ["1", "", "项目经理", "", "", "", "", ""],
            ["2", "", "系统架构师", "", "", "", "", ""],
            ["3", "", "前端工程师", "", "", "", "", ""],
            ["4", "", "后端工程师", "", "", "", "", ""],
            ["5", "", "UI工程师", "", "", "", "", ""],
            ["6", "", "测试/运维人员", "", "", "", "", ""],
        ],
        column_widths=[1.0, 2.2, 2.6, 4.2, 2.0, 2.7, 3.8, 2.3],
        center_cols={0},
    )
    add_paragraph(doc, "核心人员简历表（每名核心人员可复制填写）", bold=True)
    add_table(
        doc,
        ["项目", "填写内容"],
        [
            ["姓名", ""],
            ["拟任岗位", ""],
            ["学历/专业", ""],
            ["从业年限", ""],
            ["主要技能", ""],
            ["类似项目经验", ""],
            ["相关证书或证明材料", ""],
            ["本项目职责", ""],
            ["是否驻场及驻场周期", ""],
        ],
        column_widths=[4.8, 12.8],
    )
    add_paragraph(doc, "驻场安排及稳定性保障表", bold=True)
    add_table(
        doc,
        ["序号", "驻场岗位", "人数", "预计进场时间", "驻场地点", "驻场周期", "替补安排", "稳定性保障措施"],
        [
            ["1", "项目经理", "", "", "", "", "", ""],
            ["2", "前端工程师", "", "", "", "", "", ""],
            ["3", "后端工程师", "", "", "", "", "", ""],
            ["4", "测试/运维支持", "", "", "", "", "", ""],
        ],
        column_widths=[1.0, 2.8, 1.2, 2.4, 2.3, 2.1, 3.4, 4.7],
        center_cols={0, 2},
    )

    add_paragraph(doc, "十二、运维支持、响应保障及知识转移方案", bold=True)
    add_paragraph(doc, "服务响应及应急处置承诺表", bold=True)
    add_table(
        doc,
        ["序号", "服务事项", "响应时限", "修复/处理时限", "升级机制", "临时替代方案", "备注"],
        [
            ["1", "一般问题/咨询", "", "", "", "", ""],
            ["2", "影响业务使用的问题", "", "", "", "", ""],
            ["3", "重大问题/系统瘫痪", "", "", "", "", ""],
            ["4", "周末及节假日现网问题", "", "", "", "", ""],
        ],
        column_widths=[1.0, 3.6, 2.0, 2.4, 3.2, 4.0, 2.6],
        center_cols={0},
    )
    add_paragraph(doc, "服务治理与持续优化机制表", bold=True)
    add_table(
        doc,
        ["机制名称", "频率/触发条件", "输出物", "责任角色", "跟踪方式"],
        [
            ["项目例会机制", "", "", "", ""],
            ["周报/月报机制", "", "", "", ""],
            ["重大问题升级机制", "", "", "", ""],
            ["阶段复盘与持续优化机制", "", "", "", ""],
            ["服务考核与改进闭环机制", "", "", "", ""],
        ],
        column_widths=[4.2, 3.5, 4.5, 3.2, 3.6],
    )
    add_paragraph(doc, "运维知识库及问题闭环管理表", bold=True)
    add_table(
        doc,
        ["序号", "管理内容", "拟提交/维护材料", "更新频率", "责任人", "备注"],
        [
            ["1", "系统架构及部署说明", "", "", "", ""],
            ["2", "环境配置及数据库说明", "", "", "", ""],
            ["3", "接口清单及日志说明", "", "", "", ""],
            ["4", "常见问题及处理SOP", "", "", "", ""],
            ["5", "Bug台账及整改闭环记录", "", "", "", ""],
            ["6", "版本发布记录", "", "", "", ""],
            ["7", "巡检报告", "", "", "", ""],
        ],
        column_widths=[1.0, 4.2, 5.0, 2.5, 2.6, 3.4],
        center_cols={0},
    )
    add_paragraph(doc, "培训及知识转移计划表", bold=True)
    add_table(
        doc,
        ["序号", "培训/移交主题", "对象", "计划时间", "培训方式", "拟提交材料", "确认方式"],
        [
            ["1", "用户操作培训", "", "", "", "", ""],
            ["2", "运维管理培训", "", "", "", "", ""],
            ["3", "系统部署及配置说明", "", "", "", "", ""],
            ["4", "问题处理及应急流程", "", "", "", "", ""],
            ["5", "资料移交及答疑", "", "", "", "", ""],
        ],
        column_widths=[1.0, 4.2, 2.0, 2.2, 2.2, 4.3, 3.6],
        center_cols={0},
    )

    add_paragraph(doc, "十三、商务响应及专项承诺文件", bold=True)
    add_paragraph(doc, "采购单位关注事项响应表", bold=True)
    concern_rows = [
        ["1", "三方公司对项目开发的可参与程度", "", ""],
        ["2", "是否能够一起办公", "", ""],
        ["3", "需求沟通的便捷性及问题响应时间", "", ""],
        ["4", "软件公司提供的人员配置", "", ""],
        ["5", "合同签订后最快进场开发的时间及投入人员", "", ""],
        ["6", "对报表系统的使用经验及业绩", "", ""],
        ["7", "公司人员及规模", "", ""],
        ["8", "过往开发项目业绩与平台相似度", "", ""],
        ["9", "合同付款意向", "", ""],
        ["10", "技术方案完备合理性", "", ""],
        ["11", "需求交流参与积极性", "", ""],
        ["12", "软件报价", "", ""],
        ["13", "是否可以参与东软报表培训", "", ""],
        ["14", "技术保密、版权、赔偿", "", ""],
        ["15", "Bug承诺终生修复", "", ""],
        ["16", "对公司使用交付文档及人员培训", "", ""],
        ["17", "对公司技术力量的培养及技术问题答疑", "", ""],
        ["18", "能否接受甲方技术栈的要求", "", ""],
        ["19", "提供专人解决现网问题（包括周末、节假日）", "", ""],
        ["20", "是否接受软件开发合同条款", "", ""],
    ]
    add_table(
        doc,
        ["序号", "采购单位关注事项", "供应商响应/承诺", "对应页码"],
        concern_rows,
        column_widths=[1.0, 6.2, 7.0, 2.4],
        center_cols={0, 3},
    )
    add_paragraph(doc, "商务响应汇总表", bold=True)
    add_table(
        doc,
        ["序号", "响应事项", "谈判文件要求", "供应商响应/承诺", "是否偏离", "证明或说明材料页码"],
        [
            ["1", "付款条件", "", "", "□无偏离  □正偏离  □负偏离", ""],
            ["2", "交付要求", "", "", "□无偏离  □正偏离  □负偏离", ""],
            ["3", "验收要求", "", "", "□无偏离  □正偏离  □负偏离", ""],
            ["4", "驻场要求", "", "", "□无偏离  □正偏离  □负偏离", ""],
            ["5", "服务响应要求", "", "", "□无偏离  □正偏离  □负偏离", ""],
            ["6", "保密要求", "", "", "□无偏离  □正偏离  □负偏离", ""],
            ["7", "知识产权要求", "", "", "□无偏离  □正偏离  □负偏离", ""],
            ["8", "质保及售后服务", "", "", "□无偏离  □正偏离  □负偏离", ""],
        ],
        column_widths=[1.0, 2.8, 4.6, 4.6, 3.8, 2.6],
        center_cols={0},
    )
    add_paragraph(doc, "服务及驻场承诺函", bold=True)
    add_paragraph(doc, "致：南京三方化工设备监理有限公司")
    add_paragraph(
        doc,
        "我公司承诺，如我公司成交，将按照谈判文件、响应文件及合同约定配备项目人员，按期进场，满足驻场办公、需求沟通、开发实施、测试整改、上线保障、运维支持、培训答疑及资料移交等要求。因我公司原因导致人员不到位、响应不及时、资料不完整或服务不满足要求的，我公司愿按谈判文件及合同约定承担相应责任。",
    )
    add_signature_lines(
        doc,
        [
            "供应商名称（盖公章）：____________________",
            "法定代表人或委托代理人（签字）：____________________",
            "日期：______年______月______日",
        ],
    )
    add_paragraph(doc, "保密及知识产权承诺函", bold=True)
    add_paragraph(doc, "致：南京三方化工设备监理有限公司")
    add_paragraph(
        doc,
        "我公司承诺严格遵守谈判文件及合同关于保密、数据安全、知识产权、成果归属、源代码交付、第三方软件合规使用等要求。未经采购单位书面同意，我公司不向任何第三方披露、转让、使用或许可使用与本项目有关的资料、数据、代码、文档、业务信息及其他成果。因我公司违反前述承诺造成损失的，我公司愿按谈判文件及合同约定承担全部责任。",
    )
    add_signature_lines(
        doc,
        [
            "供应商名称（盖公章）：____________________",
            "法定代表人或委托代理人（签字）：____________________",
            "日期：______年______月______日",
        ],
    )

    add_paragraph(doc, "十四、商务条款偏离表", bold=True)
    add_table(
        doc,
        ["序号", "谈判文件条目号", "谈判文件商务条款", "响应文件商务条款", "偏离情况说明"],
        [["1", "", "", "", ""], ["2", "", "", "", ""], ["3", "", "", "", ""]],
        column_widths=[1.0, 2.8, 5.2, 5.2, 4.0],
        center_cols={0},
    )
    add_paragraph(doc, "如无偏离，则在表格中填写“无偏离”。")

    add_paragraph(doc, "十五、技术规格偏离表", bold=True)
    add_table(
        doc,
        ["序号", "货物/模块名称", "谈判文件条目号", "谈判技术条款", "响应文件技术条款", "偏离情况", "说明"],
        [["1", "", "", "", "", "", ""], ["2", "", "", "", "", "", ""], ["3", "", "", "", "", "", ""]],
        column_widths=[1.0, 2.7, 2.5, 4.2, 4.2, 2.0, 2.4],
        center_cols={0},
    )
    add_paragraph(doc, "如无偏离，则在表格中填写“无偏离”。")

    add_paragraph(doc, "十六、供应商认为必要的其他文件", bold=True)
    add_paragraph(doc, "格式自拟。供应商可以补充提交有助于证明其履约能力、服务能力、案例经验、研发能力和文件完整性的其他材料。")


def build_document() -> Path:
    paragraphs = extract_paragraphs(TENDER_DOC)
    ranges = locate_replace_ranges(paragraphs)
    score_tables = read_score_tables(SCOREBOOK)

    doc = Document()
    configure_document(doc)
    add_title(
        doc,
        "南京三方一体化数字信息平台优化提升项目谈判文件第三章及第七章替换稿",
        "依据《软件项目供应商评审评分表-最终版.xlsx》逐项对照整理",
    )
    add_paragraph(
        doc,
        "说明：本文件用于直接替换原谈判文件第三章“谈判办法”和第七章“响应文件格式”中的对应内容，重点解决评分标准描述不完整、响应文件清单与模板无法完全覆盖评分表评分点的问题。",
    )

    build_replacement_scope(doc, ranges)
    build_gap_analysis(doc)

    doc.add_page_break()
    build_chapter3(doc, score_tables, ranges)

    doc.add_page_break()
    build_chapter7(doc, ranges)

    doc.save(OUTPUT_DOC)
    return OUTPUT_DOC


if __name__ == "__main__":
    result = build_document()
    print(result)
