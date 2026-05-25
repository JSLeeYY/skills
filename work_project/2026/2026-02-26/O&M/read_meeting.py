
from docx import Document
import sys

doc = Document('2026.03.04会议纪要.docx')
output_lines = []

output_lines.append("=== 正文段落 ===")
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if t:
        output_lines.append(f"[P{i}] " + t)

output_lines.append("")
output_lines.append("=== 表格内容 ===")
for ti, table in enumerate(doc.tables):
    output_lines.append(f"--- 表格{ti+1} ---")
    for ri, row in enumerate(table.rows):
        cells = []
        seen = set()
        for c in row.cells:
            ct = c.text.replace('\n', ' ').strip()
            if ct and ct not in seen:
                cells.append(ct)
                seen.add(ct)
        if cells:
            output_lines.append(f"  行{ri}: " + " | ".join(cells))

full_text = "\n".join(output_lines)

with open("meeting_output.txt", "w", encoding="utf-8") as f:
    f.write(full_text)

print("Done! Total chars:", len(full_text))
