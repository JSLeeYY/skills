
from docx import Document

doc = Document('Operations_and_Management_SRS_Final_V11.md')
print("段落数:", len(doc.paragraphs))
print("表格数:", len(doc.tables))
